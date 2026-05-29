from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import requests
import re
import markdown as md_lib
from bs4 import BeautifulSoup

app = Flask(__name__)

@app.route('/health')
def health():
    return {'ok': True, 'service': 'onwinword'}

@app.route('/')
def index():
    return {'service': 'onwinword', 'endpoints': ['/health', '/generate', '/from-markdown']}

def fetch_image(url, timeout=20):
    try:
        r = requests.get(url, timeout=timeout, headers={'User-Agent': 'Mozilla/5.0'})
        if r.status_code == 200:
            return BytesIO(r.content)
    except Exception as e:
        print(f'[fetch_image] {url} failed: {e}')
    return None

def add_image_safe(doc, url, width_inches=5.5):
    img = fetch_image(url)
    if img:
        try:
            doc.add_picture(img, width=Inches(width_inches))
            return True
        except Exception as e:
            doc.add_paragraph(f'[图片渲染失败: {e}]')
    else:
        doc.add_paragraph(f'[图片下载失败: {url[:80]}]')
    return False

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json or {}
    doc = Document()
    if data.get('title'):
        doc.add_heading(data['title'], 0)
    for s in data.get('sections', []):
        t = s.get('type')
        try:
            if t == 'heading':
                doc.add_heading(s.get('text', ''), level=min(s.get('level', 1), 9))
            elif t == 'paragraph':
                doc.add_paragraph(s.get('text', ''))
            elif t == 'image':
                add_image_safe(doc, s['url'], s.get('width', 5.5))
                if s.get('caption'):
                    p = doc.add_paragraph(s['caption'])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.italic = True
                        run.font.size = Pt(9)
            elif t == 'image_compare':
                if s.get('label'):
                    doc.add_paragraph(s['label'], style='Heading 3')
                add_image_safe(doc, s['self_url'], s.get('width', 2.8))
                p = doc.add_paragraph('↑ 本品')
                if p.runs: p.runs[0].italic = True
                add_image_safe(doc, s['competitor_url'], s.get('width', 2.8))
                p = doc.add_paragraph('↑ 对手')
                if p.runs: p.runs[0].italic = True
                if s.get('caption'):
                    doc.add_paragraph(s['caption'])
            elif t == 'image_brief':
                if s.get('image_url'):
                    add_image_safe(doc, s['image_url'], s.get('width', 4.0))
                for label, key in [('现状', 'current'), ('问题', 'problem'), ('改图指令', 'edit_instruction'), ('合规风险', 'compliance_risk')]:
                    if s.get(key):
                        p = doc.add_paragraph()
                        run = p.add_run(f'{label}: ')
                        run.bold = True
                        if label == '合规风险':
                            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                        p.add_run(s[key])
            elif t == 'table':
                headers = s.get('headers', [])
                rows = s.get('rows', [])
                if headers:
                    tbl = doc.add_table(rows=1, cols=len(headers))
                    tbl.style = 'Light Grid Accent 1'
                    for i, h in enumerate(headers):
                        cell = tbl.rows[0].cells[i]
                        cell.text = str(h)
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                    for row in rows:
                        cells = tbl.add_row().cells
                        for i, c in enumerate(row[:len(headers)]):
                            cells[i].text = str(c) if c is not None else ''
            elif t == 'bullet':
                for item in s.get('items', []):
                    doc.add_paragraph(str(item), style='List Bullet')
            elif t == 'page_break':
                doc.add_page_break()
        except Exception as e:
            doc.add_paragraph(f'[段落渲染异常: {e}]')
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return send_file(out,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=data.get('filename', 'report.docx'))

@app.route('/from-markdown', methods=['POST'])
def from_markdown():
    data = request.json or {}
    doc = Document()
    if data.get('title'):
        doc.add_heading(data['title'], 0)
    md = data.get('markdown', '')
    image_width = data.get('image_width', 5.5)
    img_pattern = re.compile(r'\[IMG:([^\]]+)\]')
    html = md_lib.markdown(md, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    for elem in soup.children:
        if elem.name is None:
            continue
        name = elem.name.lower()
        text = elem.get_text(strip=False) if hasattr(elem, 'get_text') else ''
        img_match = img_pattern.search(text) if text else None
        if img_match:
            before = text[:img_match.start()].strip()
            if before:
                doc.add_paragraph(before)
            add_image_safe(doc, img_match.group(1), image_width)
            after = text[img_match.end():].strip()
            if after:
                doc.add_paragraph(after)
            continue
        if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(name[1])
            doc.add_heading(text.strip(), level=level)
        elif name == 'p':
            p = doc.add_paragraph()
            for child in elem.children:
                if getattr(child, 'name', None) == 'strong':
                    p.add_run(child.get_text()).bold = True
                elif getattr(child, 'name', None) == 'em':
                    p.add_run(child.get_text()).italic = True
                elif getattr(child, 'name', None) == 'code':
                    run = p.add_run(child.get_text())
                    run.font.name = 'Consolas'
                else:
                    p.add_run(str(child) if not hasattr(child, 'get_text') else child.get_text())
        elif name == 'ul':
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
        elif name == 'ol':
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style='List Number')
        elif name == 'blockquote':
            p = doc.add_paragraph(elem.get_text(strip=True))
            if p.runs: p.runs[0].italic = True
        elif name == 'pre':
            p = doc.add_paragraph(elem.get_text())
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        elif name == 'hr':
            doc.add_paragraph('—' * 30)
        elif name == 'table':
            rows = elem.find_all('tr')
            if not rows:
                continue
            header_cells = rows[0].find_all(['th', 'td'])
            tbl = doc.add_table(rows=1, cols=len(header_cells))
            tbl.style = 'Light Grid Accent 1'
            for i, h in enumerate(header_cells):
                cell = tbl.rows[0].cells[i]
                cell.text = h.get_text(strip=True)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for tr in rows[1:]:
                tds = tr.find_all(['td', 'th'])
                cells = tbl.add_row().cells
                for i, td in enumerate(tds[:len(header_cells)]):
                    cells[i].text = td.get_text(strip=True)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return send_file(out,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=data.get('filename', 'report.docx'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
